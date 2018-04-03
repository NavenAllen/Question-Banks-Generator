import numpy as np
import openpyxl as op
from docx import Document
import random

def handle_file(file):
	
	wb=op.load_workbook("./"+file)
	ws=wb.worksheets[0]

	n = 1 
	questions = [] 
	while(ws.cell(row=1, column=n).value !=None):
		ques_split = [] 
		r = 2 
		for q_type in range(0,5):
			chap_ques = []
			while(ws.cell(row=r, column=n).value !=None):
				chap_ques.append(ws.cell(row=r, column=n).value)
				r+=1 
			r+=1 
			ques_split.append(chap_ques) 
		questions.append(ques_split) 
		n+=1 

	chap_nos = n-1 
	return questions, chap_nos

def generateBank(file, questions, chap_nos, selected_chaps, qtypes, ques_nos):
	wb=op.load_workbook("./"+file)
	sheet=wb.worksheets[1]
	studMarksDetails = np.array([]) 
	n = 2 

	marksInfo = np.array([0, 0, 0, 0, 0]) 
	for i in range(0, chap_nos-1):
			marksInfo= np.vstack([marksInfo, [0.0, 0.0, 0.0, 0.0, 0.0]]) 

	while( isinstance(sheet.cell(row=2, column=n).value, int)):
		chapNo = sheet.cell(row=3, column = n).value
		qtype = sheet.cell(row=2, column =n).value
		if(qtype == 1 or qtype == 2):
			marksInfo[chapNo-1, qtype-1] += 1
		elif(qtype == 3):
			marksInfo[chapNo-1, qtype-1] += 2
		elif(qtype == 4):
			marksInfo[chapNo-1, qtype-1] +=3
		elif(qtype == 5):
			marksInfo[chapNo-1, qtype-1] +=5

		n+=1 
	
	r=4
	names = []

	while( isinstance(sheet.cell(row=r, column=1).value, str)):
		names.append(sheet.cell(row=r, column=1).value)
		r+=1 

	weights = [2, 2, 2, 2, 2] 

	n=2 
	r=4 

	while( isinstance(sheet.cell(row=r, column=n).value, int)):
		stud_details = np.array([0.0, 0.0, 0.0, 0.0, 0.0])
		for i in range(0, chap_nos-1):
			stud_details = np.vstack([stud_details, [0.0, 0.0, 0.0, 0.0, 0.0]]) 
		while( isinstance(sheet.cell(row=r, column=n).value, int)):
			chapNo = sheet.cell(row=3, column = n).value
			marks = sheet.cell(row=2, column =n).value
			stud_details[chapNo-1, marks-1] += sheet.cell(row=r, column=n).value 
			n+=1 
		if(studMarksDetails.size):
			studMarksDetails = np.vstack([studMarksDetails, [stud_details]]) 
		else:
			studMarksDetails = np.array([stud_details]) 
		n=2 
		r+=1 



	for i in range(0, np.size(studMarksDetails,0)):
		for j in range(0, np.size(studMarksDetails, 1)):
			for k in range(0, np.size(studMarksDetails, 2)):
				
				if( j+1 in selected_chaps and k+1 in qtypes):
					if(marksInfo[j, k]):
						studMarksDetails[i, j, k] = (2 - studMarksDetails[i, j, k]/marksInfo[j, k])*weights[k] 
					else:
						studMarksDetails[i, j, k] = 0.5*weights[k] 
				else:
					studMarksDetails[i, j, k] = 0
		marksSum = np.sum(studMarksDetails[i])
		studMarksDetails[i] = np.divide(studMarksDetails[i], marksSum/int(ques_nos))

	
	studMarksDetails = np.around(studMarksDetails,0) 
	

	document = Document()
	for i in range(0, np. size(studMarksDetails, 0)):
		document.add_paragraph("Question bank for " + names[i])
		for j in range(0, np.size(studMarksDetails, 1)):
			document.add_paragraph("Chapter "+str(j+1))
			for k in range(0, np.size(studMarksDetails, 2)):
				if(len(questions[j][k]) and int(studMarksDetails[i, j, k])):
					document.add_paragraph("Question Type: "+str(k+1))
					for number in range(int(studMarksDetails[i, j, k])):
						if(number<len(questions[j][k])):
							document.add_paragraph(str(number+1)+". "+questions[j][k][number])
			document.add_paragraph("")
		document.add_page_break()
	document.save('./media/Banks/Bank.docx')
	


